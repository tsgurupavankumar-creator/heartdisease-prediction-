"""
Patient Risk Prediction Module
Interactive real-time inference with Confidence Scoring and Explainability
Professional healthcare-grade UI with feature importance
"""

import streamlit as st
import joblib
import json
import os
import re
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px

from utils.ui_components import page_header
from utils.ml_pipeline import add_clinical_features, load_model

# Page configuration
st.set_page_config(
    page_title="Patient Risk Prediction",
    page_icon="🔬",
    layout="wide"
)


def get_base_estimator(model):
    """
    Extract the base estimator from calibrated models.
    """
    try:
        if hasattr(model, "calibrated_classifiers_"):
            return model.calibrated_classifiers_[0].estimator
        elif hasattr(model, "base_estimator_"):
            return model.base_estimator_
        else:
            return model
    except Exception:
        return model


def get_clinical_feature_label(feature):
    """Get clinical label for feature names."""
    feature_labels = {
        'age': 'Age (years)',
        'sex': 'Sex (0=Female, 1=Male)',
        'cp': 'Chest Pain Type (0-3)',
        'trestbps': 'Resting Blood Pressure (mm Hg)',
        'chol': 'Serum Cholesterol (mg/dL)',
        'fbs': 'Fasting Blood Sugar (0=<=120, 1=>120)',
        'restecg': 'Resting ECG Results (0-2)',
        'thalach': 'Maximum Heart Rate Achieved',
        'exang': 'Exercise Induced Angina (0=No, 1=Yes)',
        'oldpeak': 'ST Depression Induced by Exercise',
        'slope': 'Slope of Peak Exercise ST Segment (0-2)',
        'ca': 'Number of Major Vessels (0-4)',
        'thal': 'Thalassemia (0-3)',
        'age_risk_group': 'Age Risk Group (0=<45, 1=45-60, 2=60+)',
        'cholesterol_risk_group': 'Cholesterol Risk Group (0=<200, 1=200-240, 2=240+)',
        'blood_pressure_risk_group': 'BP Risk Group (0=<120, 1=120-140, 2=140+)',
        'max_heart_rate_ratio': 'Max Heart Rate / Predicted Max',
        'cholesterol_age_interaction': 'Cholesterol × Age Interaction',
        'bp_age_interaction': 'Blood Pressure × Age Interaction'
    }
    return feature_labels.get(feature, feature)


def create_feature_importance_chart(importance_df, top_n=10):
    """
    Create a professional feature importance chart using Plotly.
    Shows risk factors with increasing/decreasing impact.
    """
    
    # Take top N features
    top_df = importance_df.head(top_n).copy()
    
    # Determine if feature increases or decreases risk
    # For Random Forest, importance is always positive (magnitude)
    # We'll show the direction based on feature interpretation
    risk_directions = {
        'age': 'Increasing',
        'sex': 'Increasing (Male)',
        'cp': 'Increasing',
        'trestbps': 'Increasing',
        'chol': 'Increasing',
        'fbs': 'Increasing',
        'restecg': 'Increasing',
        'thalach': 'Decreasing',
        'exang': 'Increasing',
        'oldpeak': 'Increasing',
        'slope': 'Increasing (Downsloping)',
        'ca': 'Increasing',
        'thal': 'Increasing',
        'age_risk_group': 'Increasing',
        'cholesterol_risk_group': 'Increasing',
        'blood_pressure_risk_group': 'Increasing',
        'max_heart_rate_ratio': 'Decreasing',
        'cholesterol_age_interaction': 'Increasing',
        'bp_age_interaction': 'Increasing'
    }
    
    top_df['Direction'] = top_df['Feature'].apply(
        lambda x: risk_directions.get(x, 'Increasing')
    )
    
    # Create color mapping
    colors = []
    for _, row in top_df.iterrows():
        if 'Increasing' in row['Direction']:
            colors.append('#E53935')  # Red - Increases risk
        else:
            colors.append('#2E7D32')  # Green - Decreases risk
    
    # Create horizontal bar chart
    fig = go.Figure()
    
    fig.add_trace(go.Bar(
        x=top_df['Importance'],
        y=top_df['Feature_Label'],
        orientation='h',
        marker=dict(
            color=colors,
            line=dict(color='white', width=1)
        ),
        text=top_df['Importance'].round(4),
        textposition='outside',
        textfont=dict(size=11, color='#1F2937'),
        hovertemplate='<b>%{y}</b><br>Importance: %{x:.4f}<br>Direction: %{customdata}<extra></extra>',
        customdata=top_df['Direction']
    ))
    
    # Add reference line at mean importance
    mean_importance = top_df['Importance'].mean()
    fig.add_vline(
        x=mean_importance, 
        line_dash="dash", 
        line_color="gray",
        annotation_text=f"Mean: {mean_importance:.4f}",
        annotation_position="top right"
    )
    
    fig.update_layout(
        title=dict(
            text="<b>Feature Importance Analysis</b>",
            font=dict(size=20, color='#1E88E5')
        ),
        xaxis=dict(
            title="<b>Importance Score</b>",
            gridcolor='rgba(0,0,0,0.1)',
            zeroline=True,
            zerolinecolor='rgba(0,0,0,0.2)',
            tickfont=dict(size=12)
        ),
        yaxis=dict(
            title="",
            gridcolor='rgba(0,0,0,0.05)',
            tickfont=dict(size=13, color='#1F2937'),
            autorange='reversed'
        ),
        height=450,
        margin=dict(l=20, r=120, t=60, b=40),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(family="Arial, sans-serif"),
        legend=dict(
            x=1.02,
            y=0.98,
            title="Risk Direction",
            font=dict(size=12)
        )
    )
    
    # Add annotations for direction
    fig.add_annotation(
        x=0.02,
        y=0.02,
        xref="paper",
        yref="paper",
        text="🔴 Increasing Risk  |  🟢 Decreasing Risk",
        showarrow=False,
        font=dict(size=12, color='#1F2937'),
        bgcolor='rgba(255,255,255,0.8)',
        bordercolor='rgba(0,0,0,0.2)',
        borderwidth=1,
        borderpad=4
    )
    
    return fig


def create_risk_gauge(probability):
    """Create a professional risk gauge using Plotly."""
    
    # Determine color based on risk level
    if probability < 0.30:
        color = "#2E7D32"  # Green
        status = "Low Risk"
    elif probability < 0.60:
        color = "#F57C00"  # Orange
        status = "Moderate Risk"
    else:
        color = "#C62828"  # Red
        status = "High Risk"
    
    fig = go.Figure(go.Indicator(
        mode="gauge+number+delta",
        value=probability * 100,
        title=dict(
            text=f"<b>{status}</b>",
            font=dict(size=20, color=color)
        ),
        delta=dict(
            reference=50,
            increasing=dict(color="red"),
            decreasing=dict(color="green")
        ),
        gauge=dict(
            axis=dict(
                range=[0, 100],
                tickwidth=1,
                tickcolor="darkgray",
                tickfont=dict(size=12)
            ),
            bar=dict(color=color),
            bgcolor="rgba(0,0,0,0)",
            borderwidth=2,
            bordercolor="lightgray",
            steps=[
                dict(range=[0, 30], color="rgba(46, 125, 50, 0.2)"),
                dict(range=[30, 60], color="rgba(245, 124, 0, 0.2)"),
                dict(range=[60, 100], color="rgba(198, 40, 40, 0.2)")
            ],
            threshold=dict(
                line=dict(color="black", width=4),
                thickness=0.75,
                value=probability * 100
            )
        )
    ))
    
    fig.update_layout(
        height=250,
        margin=dict(l=20, r=20, t=50, b=20),
        font=dict(family="Arial, sans-serif")
    )
    
    return fig


def calculate_risk_factor_impact(input_values, feature_importance):
    """
    Calculate the actual impact of each feature on the risk score.
    """
    impacts = []
    
    for feature, importance in feature_importance:
        value = input_values.get(feature, 0)
        
        # Define thresholds for high risk values
        high_risk_thresholds = {
            'age': 60,
            'sex': 1,
            'cp': 2,
            'trestbps': 140,
            'chol': 240,
            'fbs': 1,
            'restecg': 1,
            'thalach': 130,
            'exang': 1,
            'oldpeak': 1.5,
            'slope': 1,
            'ca': 1,
            'thal': 1
        }
        
        threshold = high_risk_thresholds.get(feature, 0)
        
        # Calculate impact: if value is above threshold, it increases risk
        if value >= threshold:
            direction = "Increases Risk"
            impact = importance * (value / threshold if threshold > 0 else 1)
        else:
            direction = "Decreases Risk"
            impact = -importance * (1 - value / threshold if threshold > 0 else 0)
        
        impacts.append({
            'feature': feature,
            'value': value,
            'threshold': threshold,
            'importance': importance,
            'impact': impact,
            'direction': direction
        })
    
    # Sort by absolute impact
    impacts.sort(key=lambda x: abs(x['impact']), reverse=True)
    
    return impacts


def _extract_document_text(uploaded_file):
    """Extract text from supported uploaded clinical documents."""
    if uploaded_file is None:
        return "", None

    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()
    file_bytes = uploaded_file.getvalue()

    try:
        if extension == "pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return "", "PDF extraction needs the pypdf package. Run `pip install -r requirements.txt`."

            reader = PdfReader(uploaded_file)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text, None

        if extension == "docx":
            try:
                from docx import Document
            except ImportError:
                return "", "DOCX extraction needs the python-docx package. Run `pip install -r requirements.txt`."

            import io
            document = Document(io.BytesIO(file_bytes))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = []
            for table in document.tables:
                for row in table.rows:
                    table_text.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join([text, *table_text]), None

        if extension in {"txt", "csv"}:
            return file_bytes.decode("utf-8", errors="ignore"), None

        return "", "Unsupported file type. Upload PDF, DOCX, TXT, or CSV."
    except Exception as exc:
        return "", f"Could not read document: {exc}"


def _find_numeric_value(text, aliases):
    """Find the first numeric value near any accepted label."""
    for alias in aliases:
        pattern = rf"(?i)\b{re.escape(alias)}\b[^\d\-+]{{0,45}}([-+]?\d+(?:\.\d+)?)"
        match = re.search(pattern, text)
        if match:
            return float(match.group(1))
    return None


def _find_keyword_value(text, mapping):
    lowered = text.lower()
    for value, keywords in mapping.items():
        if any(keyword in lowered for keyword in keywords):
            return float(value)
    return None


def _clip(value, minimum, maximum):
    return float(max(minimum, min(maximum, value)))


def extract_clinical_values(text):
    """Best-effort extraction for common heart disease prediction fields."""
    values = {}
    if not text or not text.strip():
        return values

    normalized = re.sub(r"\s+", " ", text)

    numeric_specs = {
        "age": (["age", "patient age", "years old"], 20, 100),
        "trestbps": (["trestbps", "resting blood pressure", "resting bp", "blood pressure", "bp"], 80, 250),
        "chol": (["chol", "cholesterol", "serum cholesterol", "total cholesterol"], 100, 600),
        "thalach": (["thalach", "maximum heart rate", "max heart rate", "max hr", "peak heart rate"], 60, 220),
        "oldpeak": (["oldpeak", "st depression", "st-depression"], 0, 6.5),
        "cp": (["cp", "chest pain type", "chest pain"], 0, 3),
        "restecg": (["restecg", "resting ecg", "ecg"], 0, 2),
        "slope": (["slope", "st slope", "peak exercise st segment"], 0, 2),
        "ca": (["ca", "major vessels", "number of vessels", "vessels colored"], 0, 4),
        "thal": (["thal", "thalassemia"], 0, 3),
    }

    for feature, (aliases, minimum, maximum) in numeric_specs.items():
        value = _find_numeric_value(normalized, aliases)
        if value is not None:
            values[feature] = _clip(round(value, 1), minimum, maximum)

    sex = _find_keyword_value(normalized, {1: ["male", "sex: m", "gender: m"], 0: ["female", "sex: f", "gender: f"]})
    if sex is not None:
        values["sex"] = sex

    exang = _find_keyword_value(
        normalized,
        {1: ["exercise induced angina yes", "exang yes", "exercise angina yes"], 0: ["exercise induced angina no", "exang no", "exercise angina no"]},
    )
    if exang is not None:
        values["exang"] = exang

    fbs_value = _find_numeric_value(normalized, ["fbs", "fasting blood sugar", "fasting glucose"])
    if fbs_value is not None:
        values["fbs"] = 1.0 if fbs_value > 1 and fbs_value > 120 else _clip(round(fbs_value), 0, 1)

    categorical_text_rules = {
        "cp": {
            0: ["typical angina"],
            1: ["atypical angina"],
            2: ["non-anginal", "non anginal"],
            3: ["asymptomatic"],
        },
        "restecg": {
            0: ["ecg normal", "normal ecg", "resting ecg normal"],
            1: ["st-t abnormality", "st t abnormality"],
            2: ["left ventricular hypertrophy", "lvh"],
        },
        "slope": {
            0: ["upsloping"],
            1: ["flat slope", "slope flat"],
            2: ["downsloping"],
        },
        "thal": {
            0: ["thal normal", "normal thal"],
            1: ["fixed defect"],
            2: ["reversible defect"],
            3: ["not described"],
        },
    }

    for feature, mapping in categorical_text_rules.items():
        matched = _find_keyword_value(normalized, mapping)
        if matched is not None:
            values[feature] = matched

    return values


def apply_extracted_values(extracted_values):
    for feature, value in extracted_values.items():
        st.session_state[f"input_{feature}"] = float(value)


def app():
    """Main application function."""
    
    page_header(
        "Patient Risk Prediction",
        "Interactive interface for real-time inference with Confidence Scoring and Explainability"
    )
    
    # Check for trained model
    metadata_path = "models/model_metadata.json"
    
    if not os.path.exists(metadata_path):
        st.warning("⚠️ No trained model found. Please train a model first in Model Training Center.")
        return
    
    # Load model
    try:
        model, scaler, feature_names, metadata = load_model()
    except Exception as e:
        st.error(f"❌ Failed to load model: {str(e)}")
        return
    
    # Extract base estimator for feature importance
    base_model = get_base_estimator(model)
    
    # Display model info
    best_model_name = metadata.get("best_model", "Unknown")
    model_type = metadata.get("model_type", "Unknown")
    training_samples = metadata.get("training_samples", 0)
    metrics = metadata.get("metrics", {})
    
    st.markdown("### 🧠 Active Model Information")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Model", best_model_name)
    with col2:
        st.metric("Type", model_type)
    with col3:
        st.metric("Training Samples", training_samples)
    with col4:
        st.metric("ROC-AUC", f"{metrics.get('ROC-AUC', 0):.4f}")
    
    st.divider()
    
    # Clinical input form
    st.markdown("### 🩺 Patient Demographics and Clinical Indicators")
    
    base_features = ["age", "sex", "cp", "trestbps", "chol", "fbs", "restecg", 
                    "thalach", "exang", "oldpeak", "slope", "ca", "thal"]
    
    feature_descriptions = {
        'age': 'Age in years (20-100)',
        'sex': '0 = Female, 1 = Male',
        'cp': '0 = Typical Angina, 1 = Atypical Angina, 2 = Non-anginal Pain, 3 = Asymptomatic',
        'trestbps': 'Resting blood pressure in mm Hg (80-250)',
        'chol': 'Serum cholesterol in mg/dL (100-600)',
        'fbs': '0 = <=120 mg/dL, 1 = >120 mg/dL',
        'restecg': '0 = Normal, 1 = ST-T Abnormality, 2 = Left Ventricular Hypertrophy',
        'thalach': 'Maximum heart rate achieved (60-220)',
        'exang': '0 = No, 1 = Yes',
        'oldpeak': 'ST depression induced by exercise (0-6.5)',
        'slope': '0 = Upsloping, 1 = Flat, 2 = Downsloping',
        'ca': 'Number of major vessels colored (0-4)',
        'thal': '0 = Normal, 1 = Fixed Defect, 2 = Reversible Defect, 3 = Not Described'
    }
    
    default_values = {
        'age': 55.0, 'sex': 1.0, 'cp': 0.0, 'trestbps': 130.0, 'chol': 240.0,
        'fbs': 0.0, 'restecg': 0.0, 'thalach': 150.0, 'exang': 0.0, 'oldpeak': 1.0,
        'slope': 0.0, 'ca': 0.0, 'thal': 0.0
    }

    for feature, value in default_values.items():
        st.session_state.setdefault(f"input_{feature}", value)

    with st.container(border=True):
        st.markdown("#### Upload Clinical Document for Autofill")
        upload_col, status_col = st.columns([2, 1])
        with upload_col:
            uploaded_file = st.file_uploader(
                "Upload PDF, DOCX, TXT, or CSV",
                type=["pdf", "docx", "txt", "csv"],
                label_visibility="collapsed",
                help="The app scans the document for matching clinical fields and keeps the form editable."
            )
        with status_col:
            st.caption("Manual entry remains available. Extracted values only update matching fields.")

        if uploaded_file is not None and st.session_state.get("last_uploaded_patient_doc") != uploaded_file.name:
            document_text, extraction_error = _extract_document_text(uploaded_file)
            if extraction_error:
                st.warning(extraction_error)
            else:
                extracted_values = extract_clinical_values(document_text)
                if extracted_values:
                    apply_extracted_values(extracted_values)
                    st.session_state["last_extracted_values"] = extracted_values
                    st.success(f"Autofilled {len(extracted_values)} field(s) from {uploaded_file.name}. Review before predicting.")
                else:
                    st.info("No matching clinical values were found. You can still enter values manually.")
            st.session_state["last_uploaded_patient_doc"] = uploaded_file.name

        extracted_values = st.session_state.get("last_extracted_values", {})
        if extracted_values:
            extracted_table = pd.DataFrame(
                [
                    {
                        "Field": get_clinical_feature_label(feature),
                        "Value": value,
                    }
                    for feature, value in extracted_values.items()
                ]
            )
            st.dataframe(extracted_table, hide_index=True, use_container_width=True)
    
    with st.form("prediction_form"):
        inputs = {}
        cols = st.columns(3)
        
        for idx, feature in enumerate(base_features):
            col = cols[idx % 3]
            description = feature_descriptions.get(feature, "")
            default_val = default_values.get(feature, 0.0)
            
            inputs[feature] = col.number_input(
                feature.upper(),
                value=float(st.session_state.get(f"input_{feature}", default_val)),
                help=description,
                format="%.1f" if feature in ['age', 'trestbps', 'chol', 'thalach', 'oldpeak'] else "%.0f",
                key=f"input_{feature}"
            )
        
        submit_button = st.form_submit_button("🔍 Predict Risk", type="primary", use_container_width=True)
    
    if not submit_button:
        return
    
    # Prepare input
    input_df = pd.DataFrame([inputs])
    input_df_feat = add_clinical_features(input_df)
    
    try:
        input_final = input_df_feat[feature_names]
    except Exception as e:
        st.error(f"❌ Feature mismatch: {str(e)}")
        return
    
    # Scale and predict
    input_scaled = scaler.transform(input_final)
    prediction = model.predict(input_scaled)[0]
    probability = model.predict_proba(input_scaled)[0][1]
    
    # Confidence scoring
    confidence_margin = abs(probability - 0.5)
    if confidence_margin > 0.30:
        confidence = "High Confidence"
        confidence_icon = "✅"
    elif confidence_margin > 0.15:
        confidence = "Medium Confidence"
        confidence_icon = "ℹ️"
    else:
        confidence = "Low Confidence"
        confidence_icon = "⚠️"
    
    # Display results
    st.divider()
    st.subheader("📊 Prediction Results")
    
    # Risk classification
    if probability > 0.70:
        risk_level = "High Risk"
        risk_icon = "🔴"
        result_func = st.error
    elif probability > 0.30:
        risk_level = "Moderate Risk"
        risk_icon = "🟡"
        result_func = st.warning
    else:
        risk_level = "Low Risk"
        risk_icon = "🟢"
        result_func = st.success
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        result_func(f"{risk_icon} **{risk_level} of Heart Disease Detected**" if risk_level != "Low Risk" else f"{risk_icon} **{risk_level} of Heart Disease**")
        
        st.metric("📈 Risk Probability", f"{probability * 100:.1f}%")
        
        if risk_level == "High Risk":
            st.error(f"**Risk Level:** {risk_level}")
        elif risk_level == "Moderate Risk":
            st.warning(f"**Risk Level:** {risk_level}")
        else:
            st.success(f"**Risk Level:** {risk_level}")
        
        st.metric("🎯 Model Confidence", f"{confidence_icon} {confidence}")
        
        if confidence == "Low Confidence":
            st.warning("⚠️ Prediction uncertainty is elevated. Further evaluation strongly recommended.")
        elif confidence == "Medium Confidence":
            st.info("ℹ️ Moderate confidence. Consider additional clinical context.")
        else:
            st.success("✅ High confidence prediction. Proceed with clinical decision.")
    
    with col2:
        fig_gauge = create_risk_gauge(probability)
        st.plotly_chart(fig_gauge, use_container_width=True)
    
    # Clinical interpretation
    st.divider()
    st.subheader("🏥 Clinical Interpretation")
    
    if probability > 0.70:
        st.error("""
        🔴 **HIGH RISK**: The model predicts elevated cardiovascular risk.
        
        **Key Recommendations:**
        • Consult a cardiologist immediately
        • Perform additional diagnostic testing
        • Review cholesterol and blood pressure levels
        • Assess exercise tolerance and ECG findings
        • Consider coronary angiography if indicated
        """)
    elif probability > 0.30:
        st.warning("""
        🟡 **MODERATE RISK**: The model predicts moderate cardiovascular risk.
        
        **Key Recommendations:**
        • Schedule comprehensive cardiac evaluation
        • Monitor blood pressure and cholesterol
        • Consider stress testing or echocardiogram
        • Review lifestyle factors (diet, exercise, smoking)
        • Follow-up in 3-6 months
        """)
    else:
        st.success("""
        🟢 **LOW RISK**: The model predicts relatively low cardiovascular risk.
        
        **Key Recommendations:**
        • Continue healthy lifestyle
        • Routine screening as recommended
        • Maintain regular exercise
        • Monitor diet and weight
        • Annual check-ups
        """)
    
    # Feature importance - FIXED VERSION
    st.divider()
    st.subheader("🔍 Patient-Specific Risk Factors")
    
    try:
        # Get feature importance from base model
        if hasattr(base_model, "feature_importances_"):
            importance_values = base_model.feature_importances_
        elif hasattr(base_model, "coef_"):
            importance_values = np.abs(base_model.coef_[0])
        else:
            importance_values = None
        
        if importance_values is not None and len(importance_values) == len(feature_names):
            # Create importance DataFrame
            importance_df = pd.DataFrame({
                'Feature': feature_names,
                'Importance': importance_values
            })
            
            # Add clinical labels
            importance_df['Feature_Label'] = importance_df['Feature'].apply(get_clinical_feature_label)
            
            # Sort by importance
            importance_df = importance_df.sort_values('Importance', ascending=False)
            
            # --- DISPLAY GRAPH ---
            fig = create_feature_importance_chart(importance_df, top_n=10)
            st.plotly_chart(fig, use_container_width=True)
            
            # --- TOP RISK DRIVERS WITH IMPACT ---
            st.markdown("### 🎯 Top Risk Drivers")
            
            # Calculate actual impact for this patient
            feature_importance_list = list(zip(importance_df['Feature'], importance_df['Importance']))
            impacts = calculate_risk_factor_impact(inputs, feature_importance_list)
            
            # Display top 5 risk drivers
            for item in impacts[:5]:
                feature = item['feature']
                value = item['value']
                threshold = item['threshold']
                importance = item['importance']
                direction = item['direction']
                label = get_clinical_feature_label(feature)
                
                # Determine arrow and color
                if "Increases" in direction:
                    arrow = "⬆️"
                    color = "red"
                    impact_text = "🔴 Increases Risk"
                else:
                    arrow = "⬇️"
                    color = "green"
                    impact_text = "🟢 Decreases Risk"
                
                # Display as a card
                col1, col2, col3, col4 = st.columns([2, 1, 1, 1.5])
                with col1:
                    st.write(f"**{label}**")
                with col2:
                    st.write(f"Value: {value:.1f}")
                with col3:
                    st.write(f"Threshold: {threshold}")
                with col4:
                    if "Increases" in direction:
                        st.write(f"🔴 {impact_text}")
                    else:
                        st.write(f"🟢 {impact_text}")
                
                # Show progress bar for impact
                impact_percentage = min(abs(item['impact']) * 5, 1.0)  # Scale for display
                color_code = "#E53935" if "Increases" in direction else "#2E7D32"
                st.progress(
                    impact_percentage,
                    text=f"Impact Level: {impact_percentage*100:.0f}%"
                )
                st.caption(f"Importance Score: {importance:.4f}")
            
            # --- Detailed Impact Table ---
            with st.expander("📊 View All Feature Impacts"):
                impact_data = []
                for item in impacts:
                    label = get_clinical_feature_label(item['feature'])
                    impact_data.append({
                        'Feature': label,
                        'Value': f"{item['value']:.1f}",
                        'Threshold': f"{item['threshold']}",
                        'Direction': item['direction'],
                        'Impact Score': f"{abs(item['impact']):.4f}",
                        'Importance': f"{item['importance']:.4f}"
                    })
                
                impact_df = pd.DataFrame(impact_data)
                st.dataframe(impact_df, use_container_width=True)
            
        else:
            st.warning("Feature importance not available for this model type.")
            
    except Exception as e:
        st.warning(f"⚠️ Feature importance could not be displayed: {str(e)}")
        st.info("💡 The model is still working correctly for predictions.")
    
    # Disclaimer
    st.divider()
    st.caption("⚠️ **Medical Disclaimer:** This tool is for educational purposes only. Not a substitute for professional medical diagnosis.")


if __name__ == "__main__":
    app()
